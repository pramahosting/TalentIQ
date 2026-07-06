"""
TalentIQ - CVAnalysis Router
Resume vs Job Description ATS analyser.
Supports PDF, DOCX, TXT for both resume and job description.
"""
import io
import re
import json
import asyncio
from typing import Optional
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from db.database import get_db
from models.models import User, UserAPIKey, CVAnalysisRecord
from utils.auth_utils import get_current_user
from utils.credentials import get_credential, get_groq_model, get_all_credentials, ollama_enabled, DEFAULT_GROQ_MODEL
from utils.sequencing import next_sequence_number

router = APIRouter()


def _extract_text(content: bytes, filename: str) -> str:
    """Extract plain text from file bytes. Tries multiple libraries with fallbacks."""
    fname = (filename or "").lower().strip()

    # ── TXT ──────────────────────────────────────────────────────────────
    if fname.endswith(".txt"):
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                return content.decode(enc)
            except Exception:
                continue
        return ""

    # ── PDF ──────────────────────────────────────────────────────────────
    if fname.endswith(".pdf"):
        # Try pdfplumber first
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
                text = "\n".join(pages).strip()
                if text:
                    return text
        except Exception:
            pass

        # Try pypdf
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            pages = [p.extract_text() or "" for p in reader.pages]
            text = "\n".join(pages).strip()
            if text:
                return text
        except Exception:
            pass

        # Try PyMuPDF (fitz)
        try:
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            text = "\n".join(page.get_text() for page in doc).strip()
            if text:
                return text
        except Exception:
            pass

        return ""

    # ── DOCX ─────────────────────────────────────────────────────────────
    if fname.endswith((".docx", ".doc")):
        # Try python-docx
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))

            # Headers can contain name/email/phone in letterhead-style resumes.
            # python-docx's section.header only reads ONE header per section,
            # but a docx can store up to 3 (header1/2/3.xml) — read all via XML.
            header_footer_parts = []
            try:
                import re as _re2
                import zipfile as _zipfile
                with _zipfile.ZipFile(io.BytesIO(content)) as z:
                    for name in z.namelist():
                        if _re2.match(r"word/(header|footer)\d*\.xml$", name):
                            xml = z.read(name).decode("utf-8", errors="ignore")
                            texts = _re2.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml)
                            joined = "".join(texts).strip()
                            if joined:
                                header_footer_parts.append(joined)
            except Exception:
                pass

            paragraphs = list(header_footer_parts)
            paragraphs += [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            text = "\n".join(paragraphs).strip()
            if text:
                return text
        except Exception:
            pass

        # Try docx2txt
        try:
            import docx2txt
            text = docx2txt.process(io.BytesIO(content))
            if text and text.strip():
                return text.strip()
        except Exception:
            pass

        return ""

    # ── Fallback: try decoding as text ────────────────────────────────────
    for enc in ("utf-8", "latin-1"):
        try:
            text = content.decode(enc).strip()
            if text:
                return text
        except Exception:
            continue
    return ""


# ── KEYWORD SCORING ────────────────────────────────────────────────────────────

DOMAIN_SKILLS = [
    "python","javascript","typescript","react","node","sql","postgresql","mongodb",
    "aws","azure","gcp","docker","kubernetes","git","agile","rest","api","graphql",
    "machine learning","ai","artificial intelligence","data science","excel","power bi","tableau","salesforce",
    "figma","django","flask","java","c#","c++","go","spark","kafka","airflow","dbt",
    "snowflake","databricks","redshift","bigquery","data architecture","data governance",
    "data modelling","data modeling","data warehouse","data lake","data lakehouse","lakehouse",
    "etl","elt","data pipeline","solution design","data mesh","data fabric","data vault",
    "dimensional modelling","dimensional modeling","enterprise data warehouse","edw",
    "master data management","mdm","data quality","data catalog","data cataloguing",
    "collibra","alation","informatica","talend","teradata","hadoop","hive","hbase",
    "adls","synapse","azure data factory","azure synapse","event-driven architecture",
    "real-time data","streaming","microservices architecture","enterprise architecture",
    "solution architecture","cloud architecture","togaf","zachman",
    "stakeholder management","cloud architecture","microservices","devops","ci/cd",
    "xero","myob","quickbooks","sap","oracle","dynamics","netsuite","sage",
    "cpa","ca","acca","cma","mba","cfa","fcpa","aca","phd",
    "accounting","tax","audit","payroll","bookkeeping","bas","gst",
    "financial reporting","budgeting","forecasting","reconciliation",
    "accounts payable","accounts receivable","ifrs","gaap",
    "leadership","communication","problem solving","teamwork","stakeholder",
    "management","strategy","operations","project management","scrum","nlp","llm",
    "basel","basel iii","banking","bfsi","insurance","lending","regulatory compliance",
    "risk management","governance framework","data governance framework",
]

STOPWORDS = {"a","an","the","and","or","of","in","on","for","with","to","be","is",
             "are","it","at","from","as","by","that","this","we","you","have","has"}


# ══════════════════════════════════════════════════════════════════════════════
# ATS SCORING — two-stage structured extraction + deterministic weighted scoring
#
# The old approach asked the LLM for a single overall number directly, which
# is exactly the kind of unreliable, unrepeatable scoring that makes ATS
# tools untrustworthy. This mirrors what real ATS-checker products (Jobscan,
# Enhancv, etc.) actually do instead:
#   1. Extract structured JD requirements (hard skills, experience, education)
#   2. Extract a structured candidate profile from the resume
#   3. Compute the score with plain, transparent, reproducible Python math —
#      the LLM's job is understanding text, not doing arithmetic.
# ══════════════════════════════════════════════════════════════════════════════

def _normalize_skill(s: str) -> str:
    s = re.sub(r"\s+", " ", s.strip().lower())
    # UK/US spelling normalization — general-purpose, not a per-skill fix.
    # Without this, "dimensional modelling" and "data modeling" don't even
    # get a chance to match on their shared root word.
    for pattern, repl in _UK_TO_US_SPELLING:
        s = re.sub(pattern, repl, s)
    return s


_UK_TO_US_SPELLING = [
    (r"\bmodelling\b", "modeling"), (r"\blabelling\b", "labeling"),
    (r"\bcancelled\b", "canceled"), (r"\btravelling\b", "traveling"),
    (r"\borganisation", "organization"), (r"\bcolour", "color"),
    (r"\blicence", "license"), (r"\bcentre\b", "center"),
    (r"\bprogramme\b", "program"), (r"\banalyse", "analyze"),
    (r"\boptimise", "optimize"), (r"\bcategorise", "categorize"),
    (r"\bcustomise", "customize"), (r"\bfavour", "favor"),
    (r"\bbehaviour", "behavior"), (r"\bvisualise", "visualize"),
    (r"\bsummarise", "summarize"), (r"\bspecialise", "specialize"),
]

# Two kinds of entries, both one-directional (key = the general/JD-style
# term; values = things that, if found in a resume, PROVE the general term
# is satisfied):
#   - true synonyms/abbreviations (AI <-> Artificial Intelligence)
#   - specific technique -> general skill it's a form of (Dimensional
#     Modelling is A KIND OF Data Modeling, so it should count)
# The second category is deliberately curated rather than inferred by
# fuzzy string similarity — inferring "X modeling matches Y modeling" from
# string shape alone is exactly what causes false positives (e.g.
# "financial modeling" and "data modeling" share a word but are unrelated
# skills). Encoding actual verified domain relationships avoids that.
_SKILL_SYNONYMS = {
    "ai": ["artificial intelligence"], "artificial intelligence": ["ai"],
    "ml": ["machine learning"], "machine learning": ["ml",
        "regression", "classification", "neural network", "deep learning",
        "supervised learning", "unsupervised learning", "random forest",
        "gradient boosting", "xgboost", "scikit-learn", "tensorflow", "pytorch"],
    "bi": ["business intelligence"], "business intelligence": ["bi"],
    "power bi": ["powerbi", "power-bi"],
    "aws": ["amazon web services", "ec2", "s3", "redshift", "lambda",
        "aws glue", "amazon redshift", "cloudformation"],
    "amazon web services": ["aws"],
    "azure": ["microsoft azure", "azure data factory", "azure synapse",
        "azure synapse analytics", "adls", "adls gen2", "azure devops"],
    "gcp": ["google cloud platform", "google cloud", "bigquery", "gcp bigquery"],
    "google cloud platform": ["gcp"],
    "api": ["apis", "application programming interface", "rest api", "restful api", "graphql"],
    "apis": ["api"],
    "etl": ["extract transform load", "extract, transform, load", "elt",
        "data pipeline", "airflow", "dbt", "informatica", "talend", "ssis"],
    "elt": ["etl"],
    "data pipeline": ["etl", "elt", "airflow", "dbt", "data pipelines"],
    "sql": ["structured query language", "t-sql", "pl/sql", "mysql", "postgresql", "postgres"],
    "ci/cd": ["ci cd", "continuous integration", "continuous deployment", "jenkins", "github actions"],
    "devops": ["dev ops"],
    "nlp": ["natural language processing"], "natural language processing": ["nlp"],
    "llm": ["large language model", "large language models", "gpt", "generative ai"],
    "data governance": ["governance framework", "data governance framework",
        "data stewardship", "data catalog", "data cataloguing", "data lineage",
        "data quality framework", "collibra", "alation"],
    "edw": ["enterprise data warehouse"], "enterprise data warehouse": ["edw"],
    "mdm": ["master data management"], "master data management": ["mdm"],
    "data mesh": ["domain-oriented data", "data domain", "data products"],
    "kpi": ["key performance indicator"],
    "ux": ["user experience"], "ui": ["user interface"],
    "qa": ["quality assurance"],
    "pm": ["project management", "project manager"],
    "hr": ["human resources"],
    "crm": ["customer relationship management", "salesforce"],
    "erp": ["enterprise resource planning", "sap", "oracle erp", "netsuite"],
    # ── Data modeling / architecture: specific technique -> general skill ──
    "data modeling": [
        "dimensional modeling", "dimensional model", "data vault",
        "data vault 2.0", "star schema", "snowflake schema",
        "entity relationship modeling", "er modeling", "erd",
        "third normal form", "3nf modeling", "kimball", "inmon",
        "fsldm", "logical data modeling", "physical data modeling",
        "conceptual data modeling", "normalization", "denormalization",
    ],
    "data modelling": ["data modeling"],  # falls through to the US-spelling key above via normalization
    "data architecture": [
        "data mesh", "data fabric", "lakehouse", "data lakehouse",
        "enterprise data warehouse", "edw", "data lake", "data warehouse",
        "solution architecture", "enterprise architecture",
    ],
    "cloud architecture": ["aws", "azure", "gcp", "multi-cloud", "hybrid cloud"],
}


def _skill_present(skill: str, candidate_skills: set, resume_lower: str) -> bool:
    """A skill counts as present if any of the following hold — designed to
    catch real-world phrasing variance rather than only an exact match:
      1. it's in the LLM-extracted candidate skill list (allowing either
         side to be a substring of the other, e.g. "python" vs "python 3")
      2. the exact phrase appears literally in the resume text (after
         UK/US spelling normalization on both sides)
      3. a known synonym, abbreviation, or specific-technique-that-implies-
         the-general-skill appears in the resume text (curated list, not
         blind fuzzy matching — see _SKILL_SYNONYMS above for why)
      4. for multi-word skills, all of its significant words appear
         somewhere in the resume (not necessarily contiguous) — catches
         cases like a JD's "data governance" matching a resume's
         "established governance frameworks across enterprise data"
    Exact substring matching alone was producing false-negative "gaps" for
    skills that were genuinely present in the resume, just phrased,
    abbreviated, or spelled differently than the JD's exact wording.
    """
    sk = _normalize_skill(skill)

    if any(sk in cs or cs in sk for cs in candidate_skills):
        return True

    if sk in resume_lower:
        return True

    for variant in _SKILL_SYNONYMS.get(sk, []):
        if _normalize_skill(variant) in resume_lower:
            return True

    words = [w for w in sk.split() if len(w) > 2]
    if len(words) >= 2 and all(w in resume_lower for w in words):
        return True

    return False


def _normalize_text(s: str) -> str:
    """Same UK/US spelling normalization as _normalize_skill, applied to a
    full block of text (resume/JD) rather than a single skill phrase — both
    sides of a comparison need this or spelling normalization does nothing."""
    s = s.lower()
    for pattern, repl in _UK_TO_US_SPELLING:
        s = re.sub(pattern, repl, s)
    return s


def _compute_weighted_score(jd_req: dict, strengths: dict, resume: str) -> dict:
    """Deterministic scoring math over the extracted fields, BUT the actual
    essential/good-to-have match determination now comes directly from the
    LLM's own per-item verdict (strengths["essential_matched"/"essential_missing"]
    from utils.llm_extraction.extract_candidate_strengths) rather than being
    re-derived here via string/token matching — that matching could not
    reliably judge long capability-statement requirements or requirements
    phrased differently than the resume (e.g. "Data Modeling" vs a resume
    that says "Dimensional Modeling", a specific technique that IS a form
    of it). The deterministic _skill_present check is kept only as a
    fallback for the rare case the LLM path didn't populate these fields.
    Weights: essential coverage 50%, experience-level fit 25%, education
    fit 10%, good-to-have bonus 10%, baseline content-depth allowance 5%."""
    resume_lower = _normalize_text(resume)
    candidate_skill_set = {
        _normalize_skill(s) for s in
        (strengths.get("technical_skills", []) + strengths.get("business_skills", []))
    }

    essential = [s for s in jd_req.get("essential", []) if s]
    good_to_have = [s for s in jd_req.get("good_to_have", []) if s]

    if "essential_matched" in strengths or "essential_missing" in strengths:
        matched = strengths.get("essential_matched", [])
        missing = strengths.get("essential_missing", [])
        matched_good = strengths.get("good_to_have_matched", [])
    else:
        matched = [s for s in essential if _skill_present(_normalize_skill(s), candidate_skill_set, resume_lower)]
        missing = [s for s in essential if s not in matched]
        matched_good = [s for s in good_to_have if _skill_present(_normalize_skill(s), candidate_skill_set, resume_lower)]

    skills_pct = round(len(matched) / len(essential) * 100) if essential else 70

    min_years = jd_req.get("min_years_experience") or 0
    cand_years = strengths.get("years_experience") or 0
    if min_years <= 0:
        experience_pct = 85
    else:
        experience_pct = max(20, min(100, round(cand_years / min_years * 100)))

    edu_req = (jd_req.get("education_requirement") or "").lower()
    edu_cand = (strengths.get("education") or "").lower()
    if not edu_req:
        education_pct = 90
    elif edu_cand and (edu_cand in edu_req or edu_req in edu_cand or
                        any(w in edu_cand for w in ["bachelor", "master", "phd", "degree", "diploma"] if w in edu_req)):
        education_pct = 100
    else:
        education_pct = 45

    good_to_have_bonus_pct = min(100, len(matched_good) * 25) if good_to_have else 60

    overall = (
        skills_pct * 0.50 +
        experience_pct * 0.25 +
        education_pct * 0.10 +
        good_to_have_bonus_pct * 0.10 +
        75 * 0.05
    )
    overall = max(8, min(98, round(overall)))

    return {
        "overall": overall,
        "skills_pct": skills_pct,
        "experience_pct": experience_pct,
        "education_pct": education_pct,
        "matched": matched,
        "missing": missing,
        "matched_good_to_have": matched_good,
    }


async def _score_resume(
    resume: str, jd: str, groq_key: Optional[str], groq_model: str = DEFAULT_GROQ_MODEL,
    ollama_base_url: Optional[str] = None, ollama_model: Optional[str] = None, db=None, user_id: Optional[int] = None,
) -> dict:
    from utils.llm_extraction import (
        extract_jd_requirements_categorized, extract_candidate_strengths, extract_resume_facts,
        get_taxonomy_hint, enrich_skill_taxonomy,
    )

    known_terms = await get_taxonomy_hint(db) if db is not None else []

    # JD categorization and resume-facts extraction are genuinely
    # independent of each other — the JD's requirements don't depend on
    # the resume, and the resume's skills/experience/education don't
    # depend on the JD. Only the actual MATCHING step below needs both.
    # Running these two concurrently (rather than one after the other,
    # which is what happened before) is a real, correct optimization, not
    # just chunking for its own sake — confirmed directly in production
    # logs that each one alone typically takes ~3 seconds, so running
    # them together rather than back-to-back saves roughly that much.
    #
    # Keys are resolved SEQUENTIALLY here, before any concurrent work
    # starts — resolving both from inside the concurrent gather would mean
    # two DB writes racing on the same session, which SQLAlchemy's
    # AsyncSession doesn't allow (confirmed the hard way earlier this
    # session with a real IllegalStateChangeError).
    if db is not None and user_id is not None:
        from utils.groq_pool import resolve_groq_key
        kr_jd = await resolve_groq_key(db, user_id)
        kr_resume = await resolve_groq_key(db, user_id)
        jd_key, jd_model = kr_jd["groq_key"], kr_jd["model"] or groq_model
        resume_key, resume_model = kr_resume["groq_key"], kr_resume["model"] or groq_model
    else:
        jd_key, jd_model = groq_key, groq_model
        resume_key, resume_model = groq_key, groq_model

    jd_req, resume_facts = await asyncio.gather(
        extract_jd_requirements_categorized(jd, jd_key, jd_model, ollama_base_url, ollama_model, known_terms),
        extract_resume_facts(resume, resume_key, resume_model, ollama_base_url, ollama_model),
    )
    # If the concurrent resume-facts call didn't succeed for any reason,
    # don't pass a None down as if it were real data — just omit it, and
    # extract_candidate_strengths falls back to extracting facts itself
    # as part of the matching call, exactly as it did before this change.
    from utils.llm_extraction import _mask_key_for_log
    resume_facts_preview = _mask_key_for_log(resume_key) if resume_facts else None

    strengths = await extract_candidate_strengths(
        resume, jd_req, groq_key, groq_model, ollama_base_url, ollama_model, known_terms,
        db=db, user_id=user_id,
        pre_extracted_facts=resume_facts if resume_facts else None,
    )
    # Single, easy-to-find summary line combining every key touched across
    # THIS ENTIRE analysis (JD categorization + resume-facts extraction +
    # all candidate-strengths chunks) — the line to grep for when a
    # request could plausibly have spanned multiple pool keys and you need
    # the full picture in one place rather than piecing it together from
    # several scattered lines.
    all_key_previews = sorted(set(
        ([jd_req["_groqKeyPreview"]] if jd_req.get("_groqKeyPreview") else [])
        + ([resume_facts_preview] if resume_facts_preview else [])
        + (strengths.get("_groqKeyPreviews") or [])
    ))
    print(f"  SUMMARY: CVAnalysis request used {len(all_key_previews)} distinct Groq key(s) overall: {', '.join(all_key_previews) if all_key_previews else '(none — fell back to keyword matching)'}")

    if db is not None and strengths.get("ai_powered"):
        await enrich_skill_taxonomy(db, {
            "essential": jd_req.get("essential", []),
            "good_to_have": jd_req.get("good_to_have", []),
            "technical": strengths.get("technical_skills", []),
            "business": strengths.get("business_skills", []),
            "soft": strengths.get("soft_skills", []),
        })
    scoring = _compute_weighted_score(jd_req, strengths, resume)
    ai_powered = bool(strengths.get("ai_powered", False))

    matched, missing = scoring["matched"], scoring["missing"]
    overall = scoring["overall"]

    gaps = strengths.get("gaps") or [f"Missing required skill: {s}" for s in missing[:6]]
    if not gaps:
        gaps = ["No major gaps detected against the extracted requirements"]

    suggestions = [f'Add specific, verifiable experience with "{s}" if you have it' for s in missing[:4]]
    suggestions += [
        "Quantify achievements with concrete metrics (%, $, time saved)",
        "Mirror the exact terminology used in the job description",
    ]

    summary = strengths.get("summary") or ""
    if not summary:
        exp_clause = f" and roughly {scoring['experience_pct']}% of the target experience level" if jd_req.get("min_years_experience") else ""
        summary = (
            f"Matches {scoring['skills_pct']}% of the essential requirements{exp_clause}. "
            f"Overall fit: {'Strong' if overall >= 75 else 'Moderate' if overall >= 55 else 'Needs improvement'}."
        )

    format_warnings = [
        "Use standard ATS-compatible section headers (Experience, Education, Skills)",
        "Avoid tables, columns, or text boxes — ATS parsers often can't read them",
    ]

    # Flat "strengths" list retained for any older UI code expecting one —
    # built from the richer categorized breakdown, not a separate guess.
    flat_strengths = []
    if scoring["matched"]:
        flat_strengths.append(f"Matches {len(matched)} of {len(jd_req.get('essential', []) or [1])} essential requirements")
    flat_strengths += strengths.get("significant_experience", [])[:2]
    flat_strengths += strengths.get("certifications_degrees", [])[:2]
    if not flat_strengths:
        flat_strengths = ["Resume presents a relevant professional background"]

    return {
        "overallScore": overall,
        "strengths": flat_strengths[:6],
        "gaps": gaps[:8],
        "suggestions": suggestions[:6],
        "summaryAssessment": summary,
        "formatWarnings": format_warnings[:4],
        "detailedScores": {
            "skillsMatch": scoring["skills_pct"],
            "experience": scoring["experience_pct"],
            "tools": scoring["skills_pct"],
            "domain": max(30, overall - 5),
            "softSkills": max(30, overall - 10),
            "format": min(92, overall + 8),
        },
        "matchedSkills": matched[:15],
        "missingSkills": missing[:15],
        "aiPowered": ai_powered,
        "groqModel": groq_model if ai_powered else None,
        "groqKeyPreview": ", ".join(all_key_previews) if (ai_powered and all_key_previews) else None,
        # ── Categorized strengths — the actual point of this round's change ──
        "strengthsBreakdown": {
            "essentialMatched": strengths.get("essential_matched", []),
            "technicalSkills": strengths.get("technical_skills", []),
            "businessSkills": strengths.get("business_skills", []),
            "softSkills": strengths.get("soft_skills", []),
            "significantExperience": strengths.get("significant_experience", []),
            "certificationsDegrees": strengths.get("certifications_degrees", []),
        },
        # ── Categorized JD requirements — mirrors CandidateLens's JD Summary ──
        "jdRequirements": {
            "roleTitle": jd_req.get("role", ""),
            "location": jd_req.get("location", ""),
            "company": jd_req.get("company", ""),
            "essential": jd_req.get("essential", []),
            "goodToHave": jd_req.get("good_to_have", []),
            "optional": jd_req.get("optional", []),
            "minYearsExperience": jd_req.get("min_years_experience", 0),
            "educationRequirement": jd_req.get("education_requirement", ""),
        },
        "candidateProfile": {
            "yearsExperience": strengths.get("years_experience", 0),
            "education": strengths.get("education", ""),
        },
    }


# ── ENDPOINT ───────────────────────────────────────────────────────────────────

@router.post("/analyze")
async def analyze_resume(
    job_description: str   = Form(""),
    resume_text: str       = Form(""),
    file:     Optional[UploadFile] = File(None),
    jd_file:  Optional[UploadFile] = File(None),
    current_user: User     = Depends(get_current_user),
    db: AsyncSession       = Depends(get_db),
):
    # ── Extract resume text ────────────────────────────────────────────
    final_resume = resume_text.strip()
    if file and file.filename:
        raw = await file.read()
        extracted = _extract_text(raw, file.filename)
        if extracted.strip():
            final_resume = extracted
        elif not final_resume:
            raise HTTPException(
                400,
                f"Could not extract text from '{file.filename}'. "
                "Try copy-pasting the text directly into the text box instead."
            )

    # ── Extract JD text ────────────────────────────────────────────────
    final_jd = job_description.strip()
    if jd_file and jd_file.filename:
        raw_jd = await jd_file.read()
        extracted_jd = _extract_text(raw_jd, jd_file.filename)
        if extracted_jd.strip():
            final_jd = (final_jd + "\n" + extracted_jd).strip() if final_jd else extracted_jd
        elif not final_jd:
            raise HTTPException(
                400,
                f"Could not extract text from '{jd_file.filename}'. "
                "Try copy-pasting the job description text instead."
            )

    # ── Validate ───────────────────────────────────────────────────────
    if not final_resume:
        raise HTTPException(400, "Resume is required. Upload a PDF/DOCX/TXT file or paste the text.")
    if not final_jd:
        raise HTTPException(400, "Job description is required. Upload a file or paste the text.")

    # ── Score (structured JD/requirement extraction + deterministic weighting) ──
    # groq_key here is just the fallback default if db/user_id somehow
    # aren't available below — both extraction functions resolve their own
    # key(s) from the shared pool internally when given db + user_id,
    # including per-chunk pool draws for genuine multi-key parallelism on
    # a single request (see utils/llm_extraction.py and utils/groq_pool.py).
    groq_key = await get_credential(db, current_user.id, "groq", "api_key")
    groq_model = await get_groq_model(db, current_user.id)
    ollama_creds = await get_all_credentials(db, current_user.id, "ollama") if ollama_enabled() else {}
    ollama_base_url = ollama_creds.get("base_url")
    ollama_model = ollama_creds.get("model")

    result = await _score_resume(
        final_resume, final_jd, groq_key, groq_model,
        ollama_base_url=ollama_base_url, ollama_model=ollama_model, db=db, user_id=current_user.id,
    )

    return result


# ── HISTORY (persisted server-side, so it survives browsers/devices/refresh) ──

from pydantic import BaseModel


class SaveHistoryRequest(BaseModel):
    source_name: str = "Resume"
    overall_score: float = 0
    result: dict
    candidate_info: dict = {}
    jd_info: dict = {}


def _fmt_history(r: CVAnalysisRecord) -> dict:
    return {
        "id": r.id,
        "sequenceNumber": r.sequence_number or r.id,
        "sourceName": r.source_name,
        "overallScore": r.overall_score,
        "result": r.result or {},
        "candidateInfo": r.candidate_info or {},
        "jdInfo": r.jd_info or {},
        "createdAt": r.created_at.isoformat() if r.created_at else None,
    }


@router.post("/history")
async def save_history(
    payload: SaveHistoryRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    seq_num = await next_sequence_number(db, CVAnalysisRecord, current_user.id)
    record = CVAnalysisRecord(
        user_id=current_user.id,
        sequence_number=seq_num,
        source_name=payload.source_name,
        overall_score=payload.overall_score,
        result=payload.result,
        candidate_info=payload.candidate_info,
        jd_info=payload.jd_info,
        created_at=datetime.utcnow(),
    )
    db.add(record)
    await db.commit()
    await db.refresh(record)
    return _fmt_history(record)


@router.get("/history")
async def list_history(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(CVAnalysisRecord)
        .where(CVAnalysisRecord.user_id == current_user.id)
        .order_by(CVAnalysisRecord.created_at.desc())
        .limit(50)
    )
    return [_fmt_history(rec) for rec in r.scalars().all()]


@router.delete("/history/{record_id}")
async def delete_history_item(
    record_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(CVAnalysisRecord).where(
            CVAnalysisRecord.id == record_id,
            CVAnalysisRecord.user_id == current_user.id,
        )
    )
    rec = r.scalar_one_or_none()
    if not rec:
        raise HTTPException(404, "History item not found")
    await db.delete(rec)
    await db.commit()
    return {"message": "Deleted"}


@router.delete("/history")
async def delete_all_history(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from sqlalchemy import delete as sql_delete
    await db.execute(sql_delete(CVAnalysisRecord).where(CVAnalysisRecord.user_id == current_user.id))
    await db.commit()
    return {"message": "Deleted"}