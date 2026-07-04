"""
TalentIQ - JD Creator Router
Generates a detailed, formal Job Description (Position Description) from a
handful of inputs — role title, employment type, expiry date, required
skills/experience/education. Content is generated ENTIRELY by an LLM (Groq
first, falling back to a local/self-hosted Ollama instance) — there is no
hardcoded/templated JD text. If neither LLM is reachable, generation fails
with a clear error rather than silently returning generic content. Company
name is pulled from the user's profile (Settings), not re-entered. Produces
a downloadable .docx via python-docx.
"""
import io
import json
import requests
from datetime import datetime
from typing import List, Optional

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from pydantic import BaseModel

from db.database import get_db
from models.models import User, UserAPIKey, JDDocument
from utils.auth_utils import get_current_user
from utils.credentials import get_credential, get_all_credentials, get_groq_model, DEFAULT_GROQ_MODEL
from utils.sequencing import next_sequence_number

router = APIRouter()


# ── REQUEST / RESPONSE SCHEMAS ────────────────────────────────────────────

class JDGenerateRequest(BaseModel):
    role_title: str
    job_type: str = "Full time"          # Full time / Fix term / Contract
    contract_duration: Optional[str] = None  # e.g. "6 Months", "12 Months", or a specific end date
    expiry_date: Optional[str] = None
    skills_required: List[str] = []
    experience_required: str = ""
    education_required: str = ""


# ── AI GENERATION ────────────────────────────────────────────────────────

def _build_jd_prompt(
    role_title: str, company_name: str, job_type: str, skills: List[str],
    experience: str, education: str,
) -> str:
    return f"""You are a senior HR specialist writing a detailed, formal Position
Description document, in the style used by large organisations (e.g. universities,
corporates) — thorough, professional, and specific, not generic filler.

Role title: {role_title}
Employment type: {job_type}
Company: {company_name or "the organisation"}
Required skills: {", ".join(skills) if skills else "not specified"}
Experience required: {experience or "not specified"}
Education required: {education or "not specified"}

Write the following sections. Be specific to THIS role — infer plausible,
realistic detail from the role title, skills, and experience level given
(the way a real recruiter would write it), rather than generic statements
that could apply to any job. Do not use placeholder text.

1. "position_purpose": ONE professional paragraph (4-6 sentences) describing the
   overall purpose, scope and impact of this role within the organisation.

2. "organisational_context": ONE paragraph (3-4 sentences) describing the team or
   function this role sits within, and how it contributes to the organisation's
   broader goals.

3. "responsibilities": a list of EXACTLY 12 concise, action-oriented bullet
   points describing the key accountabilities of this role — specific enough
   to reflect the actual required skills and seniority level, not generic.

4. "required_qualifications": a list of EXACTLY 10 detailed bullet points
   covering required experience, technical skills, education, and soft
   skills — written the way a real job ad states mandatory requirements
   (e.g. "X+ years of experience in...", "Strong hands-on expertise in...",
   "Proven track record of...", "Excellent communication skills...").
   Incorporate the given skills/experience/education naturally within these.

5. "preferred_qualifications": a list of 4-6 "nice to have" bullet points —
   certifications, additional tools, or domain experience that would be a
   bonus but are not mandatory.

Return ONLY valid JSON in this exact format, no markdown, no commentary, no
text before or after the JSON object:
{{
  "position_purpose": "...",
  "organisational_context": "...",
  "responsibilities": ["...", "...", "...", "...", "...", "...", "...", "...", "...", "...", "...", "..."],
  "required_qualifications": ["...", "...", "...", "...", "...", "...", "...", "...", "...", "..."],
  "preferred_qualifications": ["...", "...", "...", "..."]
}}"""


def _parse_jd_json(raw: str) -> Optional[dict]:
    """Extract and validate the JSON payload from an LLM response."""
    text = raw.strip()
    text = text.replace("```json", "").replace("```", "").strip()
    # Some models wrap the JSON with leading commentary — grab from the first '{'
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1:
        text = text[start:end + 1]
    try:
        data = json.loads(text)
    except Exception:
        return None

    purpose = (data.get("position_purpose") or "").strip()
    context = (data.get("organisational_context") or "").strip()
    responsibilities = [r for r in data.get("responsibilities", []) if r]
    required_q = [r for r in data.get("required_qualifications", []) if r]
    preferred_q = [r for r in data.get("preferred_qualifications", []) if r]

    if not (purpose and responsibilities and required_q):
        return None

    return {
        "position_purpose": purpose,
        "organisational_context": context,
        "responsibilities": responsibilities[:15],
        "required_qualifications": required_q[:12],
        "preferred_qualifications": preferred_q[:8],
    }


def _call_groq(prompt: str, groq_key: str, groq_model: str = DEFAULT_GROQ_MODEL) -> str:
    from langchain_groq import ChatGroq
    from langchain.schema import HumanMessage

    llm = ChatGroq(api_key=groq_key, model=groq_model, temperature=0.35)
    resp = llm.invoke([HumanMessage(content=prompt)])
    return resp.content


def _call_ollama(prompt: str, base_url: str, model: str) -> str:
    url = base_url.rstrip("/") + "/api/generate"
    # Explicitly bypass any HTTP_PROXY/HTTPS_PROXY env vars for this call —
    # otherwise `requests` can route localhost/private Ollama traffic through
    # a system proxy (which 404s it), even though curl reaches it directly.
    resp = requests.post(
        url,
        json={"model": model, "prompt": prompt, "stream": False, "format": "json"},
        timeout=120,
        proxies={"http": None, "https": None},
    )
    resp.raise_for_status()
    data = resp.json()
    return data.get("response", "")


async def _generate_jd_content(
    role_title: str, company_name: str, job_type: str, skills: List[str],
    experience: str, education: str,
    groq_key: Optional[str], ollama_base_url: Optional[str], ollama_model: Optional[str],
    groq_model: str = DEFAULT_GROQ_MODEL,
) -> dict:
    """Generates JD content using an LLM only — Groq first, then Ollama.
    Raises HTTPException if no LLM is available or both fail."""
    prompt = _build_jd_prompt(role_title, company_name, job_type, skills, experience, education)

    errors = []

    if groq_key:
        try:
            raw = _call_groq(prompt, groq_key, groq_model)
            parsed = _parse_jd_json(raw)
            if parsed:
                parsed["ai_powered"] = True
                parsed["llm_provider"] = "groq"
                return parsed
            errors.append("Groq returned an unparseable response.")
        except Exception as e:
            errors.append(f"Groq error: {str(e)[:150]}")

    ollama_url = ollama_base_url or "http://localhost:11434"
    ollama_mdl = ollama_model or "llama3"
    try:
        raw = _call_ollama(prompt, ollama_url, ollama_mdl)
        parsed = _parse_jd_json(raw)
        if parsed:
            parsed["ai_powered"] = True
            parsed["llm_provider"] = "ollama"
            return parsed
        errors.append("Ollama returned an unparseable response.")
    except Exception as e:
        errors.append(f"Ollama ({ollama_url}) error: {str(e)[:150]}")

    raise HTTPException(
        400,
        "Could not generate the JD — no LLM is available. Add a Groq API key in "
        "Settings > API Keys (service: groq), or run Ollama locally (or point "
        "Settings > API Keys, service: ollama, key: base_url, at a reachable "
        f"instance) with a model pulled (e.g. `ollama pull llama3`). Details: {'; '.join(errors)}"
    )


def _fmt(d: JDDocument) -> dict:
    return {
        "id": d.id,
        "sequence_number": d.sequence_number or d.id,
        "role_title": d.role_title,
        "company_name": d.company_name,
        "job_type": d.job_type,
        "contract_duration": d.contract_duration,
        "issue_date": d.issue_date,
        "expiry_date": d.expiry_date,
        "skills_required": d.skills_required or [],
        "experience_required": d.experience_required or "",
        "education_required": d.education_required or "",
        "position_purpose": d.position_purpose or "",
        "organisational_context": d.organisational_context or "",
        "responsibilities": d.responsibilities or [],
        "required_qualifications": d.required_qualifications or [],
        "preferred_qualifications": d.preferred_qualifications or [],
        "ai_powered": d.ai_powered,
        "llm_provider": d.llm_provider or "",
        "created_at": d.created_at.isoformat() if d.created_at else None,
    }


# ── DOCX BUILDER ──────────────────────────────────────────────────────────

def _build_docx(d: JDDocument) -> bytes:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = docx.Document()

    # Base font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    def set_cell_shading(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    # ── Header ──────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(d.company_name or "Company Name")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub.add_run("Position Description")
    srun.font.size = Pt(13)
    srun.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    # ── Position Details table ───────────────────────────────────────────
    rows = 4
    include_duration = d.job_type and d.job_type.lower() != "full time" and d.contract_duration
    if include_duration:
        rows = 5
    table = doc.add_table(rows=rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    labels = ["Position Title", "Employment Type", "Date Issued", "Application Deadline"]
    values = [d.role_title, d.job_type or "Full time", d.issue_date or "", d.expiry_date or ""]
    if include_duration:
        labels.insert(2, "Contract Duration")
        values.insert(2, d.contract_duration)
    for i, (label, value) in enumerate(zip(labels, values)):
        c0, c1 = table.rows[i].cells
        c0.width = Inches(2.0)
        c1.width = Inches(4.5)
        set_cell_shading(c0, "F3F4F6")
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.size = Pt(10.5)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(str(value))
        r1.font.size = Pt(10.5)

    doc.add_paragraph()

    def add_section_heading(text):
        h = doc.add_paragraph()
        hr = h.add_run(text.upper())
        hr.bold = True
        hr.font.size = Pt(12)
        hr.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)
        pPr = h.paragraph_format
        pPr.space_before = Pt(14)
        pPr.space_after = Pt(6)
        pborder = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "0D9488")
        pborder.append(bottom)
        h._p.get_or_add_pPr().append(pborder)
        return h

    # ── Position Purpose ──────────────────────────────────────────────────
    add_section_heading("Position Purpose")
    p = doc.add_paragraph(d.position_purpose or "")
    p.paragraph_format.space_after = Pt(8)

    # ── Organisational Context ─────────────────────────────────────────────
    if d.organisational_context:
        add_section_heading("Organisational Context")
        p2 = doc.add_paragraph(d.organisational_context)
        p2.paragraph_format.space_after = Pt(8)

    # ── Key Responsibilities ───────────────────────────────────────────────
    add_section_heading("Key Responsibilities")
    for item in (d.responsibilities or []):
        bp = doc.add_paragraph(item, style="List Bullet")
        bp.paragraph_format.space_after = Pt(4)

    # ── Knowledge, Skills and Experience ────────────────────────────────────
    add_section_heading("Required Qualifications")
    if d.required_qualifications:
        for item in d.required_qualifications:
            bp = doc.add_paragraph(item, style="List Bullet")
            bp.paragraph_format.space_after = Pt(4)
    else:
        doc.add_paragraph("Not specified.")

    if d.preferred_qualifications:
        add_section_heading("Preferred Qualifications")
        for item in d.preferred_qualifications:
            bp = doc.add_paragraph(item, style="List Bullet")
            bp.paragraph_format.space_after = Pt(4)

    # ── Required Skills (raw list as entered) ───────────────────────────────
    add_section_heading("Skills Summary")
    if d.skills_required:
        for s in d.skills_required:
            bp = doc.add_paragraph(s, style="List Bullet")
            bp.paragraph_format.space_after = Pt(4)
    else:
        doc.add_paragraph("Not specified.")

    # ── Footer ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = footer.add_run(f"Generated by TalentIQ · {datetime.utcnow().strftime('%d %b %Y')}")
    frun.font.size = Pt(8)
    frun.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════════════════════
# ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/generate")
async def generate_jd(
    payload: JDGenerateRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    if not payload.role_title.strip():
        raise HTTPException(400, "Role title is required.")

    groq_key = await get_credential(db, current_user.id, "groq", "api_key")

    ollama_keys = await get_all_credentials(db, current_user.id, "ollama")
    ollama_base_url = ollama_keys.get("base_url")
    ollama_model = ollama_keys.get("model")
    groq_model = await get_groq_model(db, current_user.id)

    content = await _generate_jd_content(
        payload.role_title, current_user.company or "", payload.job_type, payload.skills_required,
        payload.experience_required, payload.education_required, groq_key, ollama_base_url, ollama_model,
        groq_model,
    )

    seq_num = await next_sequence_number(db, JDDocument, current_user.id)
    doc = JDDocument(
        user_id=current_user.id,
        sequence_number=seq_num,
        role_title=payload.role_title.strip(),
        company_name=current_user.company or "",
        job_type=payload.job_type or "Full time",
        contract_duration=payload.contract_duration if (payload.job_type and payload.job_type.lower() != "full time") else None,
        issue_date=datetime.utcnow().strftime("%d/%m/%Y"),
        expiry_date=payload.expiry_date or "",
        skills_required=payload.skills_required,
        experience_required=payload.experience_required,
        education_required=payload.education_required,
        position_purpose=content["position_purpose"],
        organisational_context=content.get("organisational_context", ""),
        responsibilities=content["responsibilities"],
        required_qualifications=content.get("required_qualifications", []),
        preferred_qualifications=content.get("preferred_qualifications", []),
        ai_powered=content["ai_powered"],
        llm_provider=content.get("llm_provider", ""),
        created_at=datetime.utcnow(),
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    return _fmt(doc)


@router.get("/documents")
async def list_documents(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(JDDocument)
        .where(JDDocument.user_id == current_user.id)
        .order_by(JDDocument.created_at.desc())
    )
    return [_fmt(d) for d in r.scalars().all()]


@router.get("/documents/{doc_id}")
async def get_document(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(JDDocument).where(JDDocument.id == doc_id, JDDocument.user_id == current_user.id)
    )
    doc = r.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "JD not found")
    return _fmt(doc)


@router.get("/documents/{doc_id}/download")
async def download_document(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(JDDocument).where(JDDocument.id == doc_id, JDDocument.user_id == current_user.id)
    )
    doc = r.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "JD not found")

    file_bytes = _build_docx(doc)
    safe_name = "".join(c for c in doc.role_title if c.isalnum() or c in (" ", "_", "-")).strip().replace(" ", "_")

    return Response(
        content=file_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=JD_{safe_name or doc.id}.docx"},
    )


@router.delete("/documents/{doc_id}")
async def delete_document(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(JDDocument).where(JDDocument.id == doc_id, JDDocument.user_id == current_user.id)
    )
    doc = r.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "JD not found")
    await db.delete(doc)
    await db.commit()
    return {"message": "Deleted"}