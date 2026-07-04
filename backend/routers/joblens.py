"""
TalentIQ - CandidateLens Router
Mirrors the original JobLens scoring logic:
  1. LLM extracts skills from JD (Groq instead of Ollama)
  2. Keyword match CV text against extracted skills
  3. Bonus for degree/experience mentions
  4. Generate interview questions via LLM
  5. Generate a 10-statement resume summary via LLM
  6. Send video-interview invite emails with a candidate-facing link
All persisted to PostgreSQL (tiq_joblens_* tables).
"""
import io
import re
import os
import json
import secrets
import smtplib
import requests
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path
from typing import List, Optional
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, BackgroundTasks
from fastapi.responses import Response
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from pydantic import BaseModel

from db.database import get_db, AsyncSessionLocal
from models.models import User, UserAPIKey, JobLensSession, JobLensCandidate
from utils.auth_utils import get_current_user
from utils.credentials import get_credential, get_all_credentials, get_groq_model, DEFAULT_GROQ_MODEL
from utils.sequencing import next_sequence_number

router = APIRouter()


@router.get("/jd-options")
async def list_jd_options(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """JD Management records for the 'New Analysis' JD-selection dropdown,
    including Client Name so it can be shown alongside the title."""
    from models.models import JDRecord, Client as ClientModel
    r = await db.execute(
        select(JDRecord).where(JDRecord.user_id == current_user.id).order_by(JDRecord.created_at.desc())
    )
    jds = r.scalars().all()
    out = []
    for jd in jds:
        client_name = ""
        if jd.client_id:
            cr = await db.execute(select(ClientModel).where(ClientModel.id == jd.client_id))
            client = cr.scalar_one_or_none()
            if client:
                client_name = client.name
        out.append({"id": jd.id, "jd_title": jd.title, "client_name": client_name, "status": jd.status})
    return out


@router.get("/vendor-candidates")
async def list_vendor_candidates_for_jd(
    jd_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """TrackedCandidates (Vendor Management / Profile Management submissions)
    for a specific JD — dynamically filtered so 'New Analysis' only offers
    profiles actually relevant to a fresh analysis run: candidates already
    Shortlisted, Selected, or Offered are excluded, since re-running ATS
    scoring on a candidate who has already progressed past that stage isn't
    useful here."""
    from models.models import TrackedCandidate, Vendor as VendorModel
    EXCLUDED_STATUSES = {"Shortlisted", "Selected", "Offered"}
    r = await db.execute(
        select(TrackedCandidate).where(TrackedCandidate.jd_id == jd_id, TrackedCandidate.user_id == current_user.id)
        .order_by(TrackedCandidate.created_at.desc())
    )
    candidates = [c for c in r.scalars().all() if c.status not in EXCLUDED_STATUSES]
    out = []
    for c in candidates:
        vr = await db.execute(select(VendorModel).where(VendorModel.id == c.vendor_id))
        vendor = vr.scalar_one_or_none()
        out.append({
            "id": c.id,
            "name": c.name,
            "vendor_id": c.vendor_id,
            "vendor_name": vendor.name if vendor else "",
            "has_resume": bool(c.resume_blob),
            "status": c.status,
        })
    return out


# ── TEXT EXTRACTION ─────────────────────────────────────────────────────────

def extract_text(content: bytes, filename: str) -> str:
    fname = (filename or "").lower()
    if fname.endswith(".txt"):
        for enc in ("utf-8", "latin-1", "cp1252"):
            try: return content.decode(enc)
            except: continue
        return ""
    if fname.endswith(".pdf"):
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                t = "\n".join(p.extract_text() or "" for p in pdf.pages).strip()
                if t: return t
        except: pass
        try:
            import pypdf
            r = pypdf.PdfReader(io.BytesIO(content))
            t = "\n".join(p.extract_text() or "" for p in r.pages).strip()
            if t: return t
        except: pass
        return ""
    if fname.endswith((".docx", ".doc")):
        # Try python-docx first (works on .docx)
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))

            # Headers often contain name/email/phone (especially in resumes
            # with a banner/letterhead layout). python-docx's section.header
            # only exposes ONE header per section by default, but a docx can
            # have up to 3 per section (default/first-page/even-page) stored
            # as header1.xml/header2.xml/header3.xml — read them all via the
            # raw XML so we never miss contact details placed there.
            header_footer_parts = []
            try:
                import re as _re2
                import zipfile as _zipfile
                with _zipfile.ZipFile(io.BytesIO(content)) as z:
                    for name in z.namelist():
                        if _re2.match(r"word/(header|footer)\d*\.xml$", name):
                            xml = z.read(name).decode("utf-8", errors="ignore")
                            texts = _re2.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml)
                            joined = "".join(texts).strip()
                            if joined:
                                header_footer_parts.append(joined)
            except Exception:
                pass

            parts = list(header_footer_parts)  # put header/footer text first
            parts += [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip(): parts.append(cell.text.strip())
            t = "\n".join(parts).strip()
            if t: return t
        except: pass
        # Try docx2txt
        try:
            import docx2txt
            t = docx2txt.process(io.BytesIO(content))
            if t and t.strip(): return t.strip()
        except: pass
        # Fallback for old binary .doc (OLE2 format): extract ASCII text stream
        if fname.endswith(".doc"):
            try:
                import re as _re
                raw = content.decode("latin-1", errors="ignore")
                chunks = _re.findall(r"[\x20-\x7e\r\n\t]{3,}", raw)
                text = "\n".join(c.strip() for c in chunks if c.strip())
                # Remove common .doc binary artifacts
                text = _re.sub(r"bjbj[a-zA-Z0-9]+", "", text)
                text = _re.sub(r"WW8Num\w+", "", text)
                text = _re.sub(r'HYPERLINK\s+"[^"]+"', "", text)
                text = _re.sub(r"\\r", "\n", text)
                text = _re.sub(r"\s{4,}", "\n", text)
                text = _re.sub(r"\n{3,}", "\n\n", text).strip()
                if len(text) > 100:
                    return text
            except: pass
        return ""
    for enc in ("utf-8", "latin-1"):
        try: return content.decode(enc).strip()
        except: continue
    return ""


# ── CANDIDATE INFO EXTRACTION ────────────────────────────────────────────────

def extract_candidate_info(text: str, filename: str) -> dict:
    """Extract name, email, phone — mirrors original extractCandidateDetails."""
    # Word documents with letterhead-style headers often have run-together
    # text with no whitespace between fields (Word renders separate <w:t>
    # runs with different formatting as visually adjacent but textually
    # concatenated), e.g. "NSW 2145resume2@gmail.comLinkedIn:". Insert a
    # space at lowercase→uppercase and digit→letter transitions so regexes
    # below can find clean boundaries — this never touches genuine words.
    norm_text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
    norm_text = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', norm_text)

    # Email
    email_m = re.search(r"[a-zA-Z][\w.+-]*@[\w.-]+\.[a-zA-Z]{2,6}", norm_text)
    email = email_m.group() if email_m else ""

    # Phone — original regex: (+?\d{1,4}[\s-]?\(?\d{1,4}\)?[\s-]?\d{3,4}[\s-]?\d{3,4})
    phone_m = re.search(r"(\+?\d{1,4}[\s\-]?\(?\d{1,4}\)?[\s\-]?\d{3,4}[\s\-]?\d{3,4})", text)
    phone = phone_m.group().replace(" ", "").strip() if phone_m else ""
    # Reject cert numbers: must have spaces/dashes or start with +
    if phone and not re.search(r"[\s\-\+\(\)]", phone_m.group() if phone_m else ""):
        phone = ""

    # Name — mirrors original: first non-email, non-4digit, <=5 word line
    name = ""
    NAME_SKIP_PATTERNS = (
        r"^page\s+\d+\s*(of\s+\d+)?$",
        r"^\d+\s*$",
        r"^confidential$",
        r"^curriculum vitae$",
        r"^resume$",
    )
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    for line in lines[:25]:
        if "@" in line:
            continue
        if re.search(r"\d{4,}", line):
            continue
        if any(re.match(p, line, re.IGNORECASE) for p in NAME_SKIP_PATTERNS):
            continue
        if len(line.split()) <= 5 and len(line) > 2:
            candidate = re.sub(r"[^a-zA-Z\s\-\.]", "", line).strip()
            if candidate and len(candidate) > 2:
                name = candidate
                break

    if not name:
        name = Path(filename).stem.replace("_", " ").replace("-", " ").title()

    # ── Experience years — look for explicit "X years" mentions ───────
    exp_years = ""
    exp_matches = re.findall(
        r"(\d{1,2})\+?\s*(?:years?|yrs?)\s*(?:of\s+)?(?:experience|exp\b)",
        text, re.IGNORECASE,
    )
    if exp_matches:
        # Take the highest mentioned figure (usually the headline summary)
        exp_years = f"{max(int(y) for y in exp_matches)}+ years"
    else:
        # Fallback: count distinct 4-digit years mentioned in work history
        # to estimate a rough career span (e.g. 2009 ... 2024)
        years_found = sorted(set(int(y) for y in re.findall(r"\b(19[7-9]\d|20[0-2]\d)\b", text)))
        if len(years_found) >= 2:
            span = years_found[-1] - years_found[0]
            if 0 < span <= 45:
                exp_years = f"~{span} years"

    # ── Summary — first substantial paragraph (career objective / profile) ──
    summary = ""
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    SUMMARY_SKIP = {"resume", "curriculum vitae", "cv", "page", "contact", "references"}
    for p in paragraphs[:15]:
        clean_p = re.sub(r"\s+", " ", p).strip()
        low = clean_p.lower()
        if len(clean_p) < 80 or len(clean_p) > 700:
            continue
        if any(s in low[:30] for s in SUMMARY_SKIP):
            continue
        if "@" in clean_p or re.search(r"^\s*[\u2022\-\*]", p):
            continue
        # Looks like prose (has multiple sentences / reasonable word count)
        if len(clean_p.split()) >= 12:
            summary = clean_p[:400]
            break

    return {
        "name": name, "email": email, "phone": phone,
        "experience_years": exp_years, "summary": summary,
    }


# ── SKILL EXTRACTION FROM JD (LLM) ──────────────────────────────────────────

_PLACEHOLDER_VALUES = {
    "nil", "n/a", "na", "none", "-", "--", "tbd", "tba", "blank", "n.a.",
    "not specified", "not applicable", "unknown", "null",
}


def _clean_extracted_field(value: Optional[str]) -> str:
    """Drops obviously-blank template placeholders (a JD table cell that
    literally says "Nil" or "N/A" is not a real role/location/company —
    treating it as one was the source of the "ROLE: Nil" bug)."""
    if not value:
        return ""
    v = value.strip()
    if not v or v.lower() in _PLACEHOLDER_VALUES or len(v) > 120:
        return ""
    return v


async def extract_jd_details(jd_text: str, groq_key: str, groq_model: str = DEFAULT_GROQ_MODEL) -> dict:
    """Extracts role title, location, company, and skills (categorized into
    Essential / Good to Have / Optional) from the JD in a single LLM call —
    used for both the CandidateLens "Job Description Summary" panel and the
    skill-matching logic below."""
    try:
        from langchain_groq import ChatGroq
        from langchain.schema import HumanMessage

        llm = ChatGroq(api_key=groq_key, model=groq_model, temperature=0.1)
        prompt = f"""You are a recruitment AI assistant. Read the job description below and
extract these fields precisely. If a field genuinely isn't stated anywhere
in the text, return an empty string for it — do NOT guess, and do NOT
return placeholder text like "Nil", "N/A", or "TBD" as if it were a real
value.

Also categorize every required/desired skill or requirement into exactly
one of three tiers, based on how the JD phrases it:
- "essential": stated as required/must-have/mandatory
- "good_to_have": stated as preferred/desirable/advantageous but not mandatory
- "optional": mentioned only in passing, or a minor/bonus item

Job Description:
\"\"\"{jd_text[:3000]}\"\"\"

Return ONLY valid JSON in this format:
{{
  "role": "<job title, or empty string if not stated>",
  "location": "<work location/city, or empty string if not stated>",
  "company": "<hiring company name, or empty string if not stated>",
  "essential": ["skill1", "skill2"],
  "good_to_have": ["skill3", "skill4"],
  "optional": ["skill5"]
}}"""

        resp = llm.invoke([HumanMessage(content=prompt)])
        raw = resp.content.strip().replace("```json", "").replace("```", "").strip()
        data = json.loads(raw)
        essential = [s.lower().strip() for s in data.get("essential", []) if s]
        good_to_have = [s.lower().strip() for s in data.get("good_to_have", []) if s]
        optional = [s.lower().strip() for s in data.get("optional", []) if s]
        return {
            "role": _clean_extracted_field(data.get("role")),
            "location": _clean_extracted_field(data.get("location")),
            "company": _clean_extracted_field(data.get("company")),
            "essential": essential,
            "good_to_have": good_to_have,
            "optional": optional,
            "skills": list(dict.fromkeys(essential + good_to_have + optional)),  # flat list — existing scoring logic
        }
    except Exception:
        return _heuristic_jd_details(jd_text)


async def extract_skills_from_jd(jd_text: str, groq_key: str) -> list:
    """Kept for any other call sites that only want the skill list."""
    details = await extract_jd_details(jd_text, groq_key)
    return details["skills"]


def _keyword_extract_jd(jd_text: str) -> list:
    """Fallback keyword extraction — returns only real skills, not JD prose."""
    DOMAIN_SKILLS = [
        "python","javascript","typescript","react","node","sql","postgresql","mongodb",
        "aws","azure","gcp","docker","kubernetes","git","agile","rest","api","graphql",
        "machine learning","ai","artificial intelligence","data science","excel","power bi","tableau","salesforce",
        "figma","django","flask","java","c#","c++","go","spark","kafka","dbt","snowflake",
        "databricks","redshift","bigquery","data architecture","data governance","etl",
        "data mesh","data fabric","data vault","dimensional modelling","dimensional modeling",
        "enterprise data warehouse","edw","lakehouse","master data management","mdm",
        "data quality","data catalog","collibra","alation","informatica","talend",
        "teradata","hive","hbase","adls","synapse","azure data factory","azure synapse",
        "event-driven architecture","real-time data","streaming","enterprise architecture",
        "solution architecture","zachman","basel","basel iii","banking","bfsi","insurance",
        "lending","regulatory compliance","risk management","governance framework",
        "xero","myob","quickbooks","sap","oracle","dynamics","netsuite",
        "cpa","ca","acca","cma","mba","cfa","phd","accounting","tax","audit","payroll",
        "financial reporting","budgeting","forecasting","reconciliation","ifrs","gaap",
        "leadership","communication","problem solving","scrum","project management",
        "togaf","pmp","csm","hadoop","datastage","tibco","react native",
        "devops","ci/cd","terraform","ansible","linux","bash","swift","kotlin",
        "microservices","restful","graphql","redis","elasticsearch","rabbitmq",
    ]
    jd_lower = jd_text.lower()
    found = [s for s in DOMAIN_SKILLS if s in jd_lower]
    return list(dict.fromkeys(found))[:30]


def _heuristic_jd_details(jd_text: str) -> dict:
    """Non-LLM fallback for role/location/company — same placeholder
    filtering as the LLM path, so a JD with a literal 'Position: Nil'
    template field doesn't get treated as a real job title."""
    role_m = re.search(r"(?:job\s*title|role|position\s*title)\s*[:\-]\s*(.+)", jd_text, re.IGNORECASE)
    loc_m = re.search(r"(?:location|based\s*in|located\s*in)\s*[:\-]\s*(.+)", jd_text, re.IGNORECASE)
    comp_m = re.search(r"(?:company|organisation|employer)\s*[:\-]\s*(.+)", jd_text, re.IGNORECASE)
    skills = _keyword_extract_jd(jd_text)
    return {
        "role": _clean_extracted_field(role_m.group(1).split("\n")[0] if role_m else None),
        "location": _clean_extracted_field(loc_m.group(1).split("\n")[0] if loc_m else None),
        "company": _clean_extracted_field(comp_m.group(1).split("\n")[0] if comp_m else None),
        "essential": skills,
        "good_to_have": [],
        "optional": [],
        "skills": skills,
    }


# ── SCORING (mirrors calculateScore exactly) ─────────────────────────────────

# Same synonym/abbreviation set used in CVAnalysis and JobHunter — plain
# substring matching alone produces false-negative "gaps"/"missing skills"
# for skills genuinely present but phrased/abbreviated/spelled differently
# than the JD's exact wording (e.g. resume says "ML", JD extraction says
# "Machine Learning"; or resume says "Dimensional Modelling" — a specific
# technique that IS a form of "Data Modeling" — which used to show as a
# false-negative gap since it's not a literal substring match at all).
_UK_TO_US_SPELLING = [
    (r"\bmodelling\b", "modeling"), (r"\blabelling\b", "labeling"),
    (r"\bcancelled\b", "canceled"), (r"\btravelling\b", "traveling"),
    (r"\borganisation", "organization"), (r"\bcolour", "color"),
    (r"\blicence", "license"), (r"\bcentre\b", "center"),
    (r"\bprogramme\b", "program"), (r"\banalyse", "analyze"),
    (r"\boptimise", "optimize"), (r"\bcategorise", "categorize"),
    (r"\bcustomise", "customize"), (r"\bfavour", "favor"),
    (r"\bbehaviour", "behavior"), (r"\bvisualise", "visualize"),
    (r"\bsummarise", "summarize"), (r"\bspecialise", "specialize"),
]


def _normalize_text(s: str) -> str:
    s = s.lower()
    for pattern, repl in _UK_TO_US_SPELLING:
        s = re.sub(pattern, repl, s)
    return s


# Two kinds of entries, both one-directional (key = the general/JD-style
# term; values = things that, if found in a resume, PROVE the general term
# is satisfied): true synonyms/abbreviations, and specific-technique ->
# general-skill-it's-a-form-of (curated, not fuzzy-string-guessed, so it
# doesn't cause false positives the way blind similarity matching would).
_SKILL_SYNONYMS = {
    "ai": ["artificial intelligence"], "artificial intelligence": ["ai"],
    "ml": ["machine learning"], "machine learning": ["ml",
        "regression", "classification", "neural network", "deep learning",
        "supervised learning", "unsupervised learning", "random forest",
        "gradient boosting", "xgboost", "scikit-learn", "tensorflow", "pytorch"],
    "bi": ["business intelligence"], "business intelligence": ["bi"],
    "power bi": ["powerbi", "power-bi"],
    "aws": ["amazon web services", "ec2", "s3", "redshift", "lambda",
        "aws glue", "amazon redshift", "cloudformation"],
    "amazon web services": ["aws"],
    "azure": ["microsoft azure", "azure data factory", "azure synapse",
        "azure synapse analytics", "adls", "adls gen2", "azure devops"],
    "gcp": ["google cloud platform", "google cloud", "bigquery", "gcp bigquery"],
    "google cloud platform": ["gcp"],
    "api": ["apis", "application programming interface", "rest api", "restful api", "graphql"],
    "apis": ["api"],
    "etl": ["extract transform load", "extract, transform, load", "elt",
        "data pipeline", "airflow", "dbt", "informatica", "talend", "ssis"],
    "elt": ["etl"],
    "data pipeline": ["etl", "elt", "airflow", "dbt", "data pipelines"],
    "sql": ["structured query language", "t-sql", "pl/sql", "mysql", "postgresql", "postgres"],
    "ci/cd": ["ci cd", "continuous integration", "continuous deployment", "jenkins", "github actions"],
    "devops": ["dev ops"],
    "nlp": ["natural language processing"], "natural language processing": ["nlp"],
    "llm": ["large language model", "large language models", "gpt", "generative ai"],
    "data governance": ["governance framework", "data governance framework",
        "data stewardship", "data catalog", "data cataloguing", "data lineage",
        "data quality framework", "collibra", "alation"],
    "edw": ["enterprise data warehouse"], "enterprise data warehouse": ["edw"],
    "mdm": ["master data management"], "master data management": ["mdm"],
    "data mesh": ["domain-oriented data", "data domain", "data products"],
    "kpi": ["key performance indicator"],
    "ux": ["user experience"], "ui": ["user interface"],
    "qa": ["quality assurance"],
    "pm": ["project management", "project manager"],
    "hr": ["human resources"],
    "crm": ["customer relationship management", "salesforce"],
    "erp": ["enterprise resource planning", "sap", "oracle erp", "netsuite"],
    "data modeling": [
        "dimensional modeling", "dimensional model", "data vault",
        "data vault 2.0", "star schema", "snowflake schema",
        "entity relationship modeling", "er modeling", "erd",
        "third normal form", "3nf modeling", "kimball", "inmon",
        "fsldm", "logical data modeling", "physical data modeling",
        "conceptual data modeling", "normalization", "denormalization",
    ],
    "data architecture": [
        "data mesh", "data fabric", "lakehouse", "data lakehouse",
        "enterprise data warehouse", "edw", "data lake", "data warehouse",
        "solution architecture", "enterprise architecture",
    ],
    "cloud architecture": ["aws", "azure", "gcp", "multi-cloud", "hybrid cloud"],
}


def calculate_score(cv_text: str, jd_skills: list) -> dict:
    """Direct port of the original JS calculateScore function, with matching
    made tolerant of synonyms/abbreviations/spelling and specific-technique
    -> general-skill relationships (same rationale as CVAnalysis's
    _skill_present — exact substring alone produces false-negative "gaps"
    for skills that are genuinely present but phrased differently)."""
    # Clean CV text same way as original, then spelling-normalize
    cv_lower = _normalize_text(cv_text)
    cv_lower = re.sub(r"\b(an|the|and|or|of|in|on|for|with|to|be|is|are|it|at|from|as|by|that|this|if|we|you)\b", "", cv_lower)
    cv_lower = re.sub(r"[^a-z0-9+#.\-]", " ", cv_lower)

    def _present(skill: str) -> bool:
        sk = _normalize_text(skill)
        if sk in cv_lower:
            return True
        for variant in _SKILL_SYNONYMS.get(sk, []):
            if _normalize_text(variant) in cv_lower:
                return True
        words = [w for w in sk.split() if len(w) > 2]
        if len(words) >= 2 and all(w in cv_lower for w in words):
            return True
        return False

    matched = [s for s in jd_skills if _present(s)]
    gap     = [s for s in jd_skills if s not in matched]

    score = (len(matched) / len(jd_skills) * 100) if jd_skills else 0
    bonus = 0
    reasons = []

    if re.search(r"bachelor|master|degree", cv_lower):
        bonus += 10
        reasons.append("Degree +10")

    if re.search(r"experience|\d+\s+(years|year)", cv_lower):
        bonus += 10
        reasons.append("Experience +10")

    score = min(100, score + bonus)

    return {
        "score":   round(score, 1),
        "matched": matched,
        "gap":     gap,
        "bonus":   bonus,
        "reasons": "; ".join(reasons),
    }


# ── QUESTION GENERATION ──────────────────────────────────────────────────────

async def generate_questions(jd_text: str, candidate_name: str, matched_skills: list, groq_key: str, groq_model: str = DEFAULT_GROQ_MODEL) -> list:
    """Mirrors buildQuestionPrompt + callOllamaGenerate."""
    try:
        from langchain_groq import ChatGroq
        from langchain.schema import HumanMessage

        llm = ChatGroq(api_key=groq_key, model=groq_model, temperature=0.4)
        skills_str = ", ".join(matched_skills[:8]) if matched_skills else "relevant skills"
        prompt = f"""You are a recruitment AI assistant.
Generate exactly 5 interview questions for {candidate_name} based on the Job Description below.
Focus on required skills ({skills_str}), tools, and responsibilities.

Job Description:
\"\"\"{jd_text[:1500]}\"\"\"

Return ONLY valid JSON:
{{
  "questions": ["Question 1","Question 2","Question 3","Question 4","Question 5"]
}}"""

        resp = llm.invoke([HumanMessage(content=prompt)])
        raw = resp.content.strip().replace("```json", "").replace("```", "").strip()
        data = json.loads(raw)
        return data.get("questions", [])[:5]
    except Exception:
        return _default_questions(candidate_name, matched_skills)


def _default_questions(name: str, skills: list) -> list:
    qs = []
    if skills:
        qs.append(f"Tell me about a project where you used {skills[0]}.")
        if len(skills) > 1:
            qs.append(f"Rate your proficiency in {skills[1]} and give a real example.")
        if len(skills) > 2:
            qs.append(f"What challenges have you faced with {skills[2]}?")
    qs.append(f"Why are you the right candidate for this role, {name}?")
    qs.append("Where do you see yourself in 3 years?")
    return qs[:5]


# ── RESUME SUMMARY (10 statements) ──────────────────────────────────────────

async def generate_resume_summary(cv_text: str, groq_key: Optional[str], groq_model: str = DEFAULT_GROQ_MODEL) -> dict:
    """Produce a categorized resume summary — multiple specific bullet
    points grouped under Experience, Skills, Education, Achievements, and
    Availability & Work Rights — rather than one flat list of generic
    sentences. Each bullet should surface a genuinely relevant, specific
    detail (a role, a scale, a result, a named skill), not a filler
    restatement of the section heading. Uses Groq LLM when a key is
    available, otherwise falls back to heuristic keyword extraction."""
    if groq_key:
        try:
            from langchain_groq import ChatGroq
            from langchain.schema import HumanMessage

            llm = ChatGroq(api_key=groq_key, model=groq_model, temperature=0.2)
            prompt = f"""You are a recruitment analyst producing a sharp, specific candidate
summary for a recruiter who is short on time. Read the resume below and
extract the MOST relevant and important points — prioritize specifics
(role titles, years, scale, named technologies, quantified results) over
generic statements. Do not invent facts the resume doesn't support.

Resume:
\"\"\"{cv_text[:5000]}\"\"\"

Return ONLY valid JSON, no markdown, no commentary, in this exact format:
{{
  "experience": ["<specific, substantive bullet about a role/scope/achievement>", "..."],
  "skills": ["<specific bullet grouping related skills or naming a standout one>", "..."],
  "education": ["<specific bullet — degree, institution, year if stated>", "..."],
  "achievements": ["<specific, quantified accomplishment if the resume supports one>", "..."],
  "availability_work_rights": ["<specific bullet if stated>", "..."]
}}

Rules:
- experience: 3-5 bullets, each about a distinct role, project, or scope of responsibility — not one bullet per job title, but the most IMPORTANT/relevant parts of their experience
- skills: 2-4 bullets, grouping related skills together rather than one skill per bullet
- education: 1-2 bullets
- achievements: 1-3 bullets — only include if the resume actually states a concrete result/metric; omit entirely (empty list) rather than inventing one
- availability_work_rights: 0-2 bullets — only include if the resume actually mentions notice period, availability, citizenship, or work rights; omit entirely if not mentioned
- Every bullet must be a full, specific sentence a recruiter could act on — not a category label restated as a sentence"""

            resp = llm.invoke([HumanMessage(content=prompt)])
            raw = resp.content.strip().replace("```json", "").replace("```", "").strip()
            start, end = raw.find("{"), raw.rfind("}")
            if start != -1 and end != -1:
                raw = raw[start:end + 1]
            data = json.loads(raw)
            result = {
                "experience": [s for s in data.get("experience", []) if s][:5],
                "skills": [s for s in data.get("skills", []) if s][:4],
                "education": [s for s in data.get("education", []) if s][:2],
                "achievements": [s for s in data.get("achievements", []) if s][:3],
                "availability_work_rights": [s for s in data.get("availability_work_rights", []) if s][:2],
            }
            if any(result.values()):
                return result
        except Exception:
            pass
    return _fallback_resume_summary(cv_text)


def _fallback_resume_summary(cv_text: str) -> dict:
    """Heuristic categorized resume summary used when no Groq key is configured."""
    low = cv_text.lower()
    result: dict = {"experience": [], "skills": [], "education": [], "achievements": [], "availability_work_rights": []}

    edu_m = re.search(r"(bachelor|master|phd|doctorate|diploma|degree)[^\n.]{0,90}", low)
    if edu_m:
        result["education"].append(edu_m.group().strip().capitalize() + ".")

    keyword_bank = [
        "python", "java", "javascript", "sql", "excel", "power bi", "tableau",
        "aws", "azure", "gcp", "docker", "kubernetes", "project management",
        "accounting", "communication", "leadership", "react", "node",
        "salesforce", "sap", "financial reporting", "data analysis",
    ]
    found_skills = [s for s in keyword_bank if s in low]
    if found_skills:
        result["skills"].append(f"Resume indicates experience with {', '.join(found_skills[:6])}.")

    exp_m = re.search(r"(\d{1,2})\+?\s*(?:years?|yrs?)\s*(?:of\s+)?experience", low)
    if exp_m:
        result["experience"].append(f"Approximately {exp_m.group(1)}+ years of relevant experience.")
    if re.search(r"present|current(ly)?\s+work", low):
        result["experience"].append("Resume references a current or recent professional role.")
    if re.search(r"project", low):
        result["experience"].append("Resume highlights project-based work experience.")
    if not result["experience"]:
        result["experience"].append("Years and scope of experience not explicitly stated in resume.")

    if re.search(r"immediate(ly)?\s+available|available\s+immediately|notice\s+period", low):
        avail_m = re.search(r"([a-z0-9 ]{0,10}notice\s+period[a-z0-9 ]{0,15}|immediate(ly)?\s+available)", low)
        result["availability_work_rights"].append((avail_m.group().strip().capitalize() + ".") if avail_m else "Availability mentioned in resume.")
    if re.search(r"citizen|permanent resident|\bpr\b|work visa|work rights|unrestricted work rights|485 visa|482 visa|sponsorship", low):
        cit_m = re.search(r"([a-z0-9 ]{0,20}(citizen|permanent resident|work visa|work rights|sponsorship)[a-z0-9 ]{0,20})", low)
        result["availability_work_rights"].append((cit_m.group().strip().capitalize() + ".") if cit_m else "Citizenship/work rights mentioned in resume.")

    if re.search(r"certificat", low):
        result["achievements"].append("Resume mentions professional certification(s).")

    return result


# ── SMTP EMAIL SENDING ────────────────────────────────────────────────────────

async def _get_smtp_config(user_id: int, db: AsyncSession) -> dict:
    # SMTP is strictly private — never shared, never falls back to another
    # user's or admin's credentials.
    return await get_all_credentials(db, user_id, "smtp")


def _send_email(smtp_cfg: dict, to_email: str, subject: str, html_body: str):
    host = smtp_cfg.get("host")
    port = int(smtp_cfg.get("port") or 587)
    username = smtp_cfg.get("username")
    password = smtp_cfg.get("password")
    from_email = smtp_cfg.get("from_email") or username

    if not (host and username and password and from_email):
        raise HTTPException(
            400,
            "SMTP is not configured. Add credentials in Settings > API Keys "
            "(service: smtp; key names: host, port, username, password, from_email)."
        )

    msg = MIMEMultipart("alternative")
    msg["Subject"] = subject
    msg["From"] = from_email
    msg["To"] = to_email
    msg.attach(MIMEText(html_body, "html"))

    try:
        with smtplib.SMTP(host, port, timeout=15) as server:
            server.starttls()
            server.login(username, password)
            server.sendmail(from_email, [to_email], msg.as_string())
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Failed to send email: {str(e)[:200]}")


class SendInviteRequest(BaseModel):
    to_email: str
    subject: str
    body_html: str


# ── FORMAT CANDIDATE ─────────────────────────────────────────────────────────

def _fmt(c: JobLensCandidate) -> dict:
    return {
        "id": c.id,
        "name": c.name,
        "email": c.email,
        "phone": c.phone,
        "filename": c.filename,
        "ats_score": round(c.ats_score, 1),
        "status": c.status,
        "matched_skills": c.matched_skills or [],
        "missing_skills": c.missing_skills or [],
        "bonus": c.bonus,
        "bonus_reasons": c.bonus_reasons,
        "experience_years": c.experience_years or "",
        "summary": c.summary or "",
        "resume_summary": c.resume_summary or [],
        "interview_token": c.interview_token,
        "contacted": bool(c.contacted),
        "video_status": c.video_status,
        "shortlisted": c.shortlisted,
        "interview_questions": c.interview_questions or [],
        "emotion_happy": c.emotion_happy,
        "emotion_neutral": c.emotion_neutral,
        "emotion_sad": c.emotion_sad,
        "emotion_angry": c.emotion_angry,
        "emotion_fear": c.emotion_fear,
        "emotion_disgust": c.emotion_disgust,
        "emotion_surprise": c.emotion_surprise,
        "dominant_emotion": c.dominant_emotion,
        "has_resume_file": bool(c.resume_file_blob),
        "has_video": bool(c.video_blob),
        "video_transcript": c.video_transcript or "",
        "video_analysis": c.video_analysis or None,
        "video_analysis_status": c.video_analysis_status or "Pending",
        "source_vendor_id": c.source_vendor_id,
        "source_vendor_name": c.source_vendor_name or "",
        "strengths_breakdown": c.strengths_breakdown or None,
    }


# ══════════════════════════════════════════════════════════════════════════════
# ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/run")
async def run_joblens(
    jd_text: str = Form(""),
    low_threshold: int = Form(40),
    high_threshold: int = Form(70),
    jd_file: Optional[UploadFile] = File(None),
    cv_files: List[UploadFile] = File(default=[]),
    jd_record_id: Optional[int] = Form(None),          # NEW: pull JD from JD Management instead of text/file
    source_candidate_ids: str = Form(""),               # NEW: comma-separated TrackedCandidate ids from Vendor Management
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    # ── Extract JD — either from JD Management (new) or text/file upload (existing) ──
    final_jd = jd_text.strip()
    jd_client_name = ""
    if jd_record_id:
        from models.models import JDRecord, Client as ClientModel
        jr = await db.execute(select(JDRecord).where(JDRecord.id == jd_record_id, JDRecord.user_id == current_user.id))
        jd_record = jr.scalar_one_or_none()
        if not jd_record:
            raise HTTPException(404, "Selected JD not found.")
        # Use the JD's stored description as the JD text; title/client used for the summary panel
        final_jd = jd_record.description or jd_record.title
        jd_client_name = ""
        if jd_record.client_id:
            cr = await db.execute(select(ClientModel).where(ClientModel.id == jd_record.client_id))
            client = cr.scalar_one_or_none()
            if client:
                jd_client_name = client.name
    elif jd_file and jd_file.filename:
        raw = await jd_file.read()
        extracted = extract_text(raw, jd_file.filename)
        if extracted.strip():
            final_jd = extracted

    # ── Resolve candidate sources: uploaded files (existing) + Vendor
    #    Management submissions (new) — both can be used together.
    from models.models import TrackedCandidate, Vendor as VendorModel
    source_candidates = []
    if source_candidate_ids.strip():
        ids = [int(x) for x in source_candidate_ids.split(",") if x.strip().isdigit()]
        if ids:
            tcr = await db.execute(
                select(TrackedCandidate).where(TrackedCandidate.id.in_(ids), TrackedCandidate.user_id == current_user.id)
            )
            source_candidates = tcr.scalars().all()

    if not final_jd:
        raise HTTPException(400, "Job description is required.")
    if not cv_files and not source_candidates:
        raise HTTPException(400, "At least one CV is required (upload files, or select candidates from Vendor Management).")

    # ── Get Groq key (own key first, falls back to admin-configured global) ──
    groq_key = await get_credential(db, current_user.id, "groq", "api_key")
    groq_model = await get_groq_model(db, current_user.id)

    # ── Extract JD details: role, location, company, categorized skills (LLM or heuristic) ──
    # If this session was started from an existing JD Management record that
    # already has persisted categorized requirements, reuse them directly —
    # re-extracting the same JD's requirements via LLM on every single
    # analysis run would be wasteful and could drift from what's shown in
    # JD Management itself.
    if jd_record_id and jd_record and (jd_record.essential_skills or jd_record.good_to_have_skills):
        jd_details = {
            "role": jd_record.title,
            "location": "",
            "company": jd_client_name,
            "essential": jd_record.essential_skills or [],
            "good_to_have": jd_record.good_to_have_skills or [],
            "optional": jd_record.optional_skills or [],
            "skills": list(dict.fromkeys(
                (jd_record.essential_skills or []) + (jd_record.good_to_have_skills or []) + (jd_record.optional_skills or [])
            )),
        }
    elif groq_key:
        jd_details = await extract_jd_details(final_jd, groq_key, groq_model)
    else:
        jd_details = _heuristic_jd_details(final_jd)
    jd_skills = jd_details["skills"]

    # ── Create session ─────────────────────────────────────────────────
    try:
        seq_num = await next_sequence_number(db, JobLensSession, current_user.id)
        session = JobLensSession(
            sequence_number=seq_num,
            user_id=current_user.id,
            jd_text=final_jd,
            jd_skills=jd_skills,
            jd_role=jd_details["role"],
            jd_location=jd_details["location"],
            jd_company=jd_details["company"],
            jd_record_id=jd_record_id,
            jd_client_name=jd_client_name,
            jd_essential_skills=jd_details.get("essential", []),
            jd_good_to_have_skills=jd_details.get("good_to_have", []),
            jd_optional_skills=jd_details.get("optional", []),
            low_threshold=low_threshold,
            high_threshold=high_threshold,
            status="completed",
            cv_count=len(cv_files) + len(source_candidates),
            created_at=datetime.utcnow(),
        )
        db.add(session)
        await db.commit()
        await db.refresh(session)
    except Exception as e:
        raise HTTPException(500, f"DB error creating session: {str(e)}")

    # ── Score each CV ──────────────────────────────────────────────────
    async def _score_and_build_candidate(
        content: bytes, filename: str,
        source_vendor_id=None, source_vendor_name=None, source_tracked_candidate_id=None,
    ):
        cv_text = extract_text(content, filename)
        if not cv_text.strip():
            return None

        info   = extract_candidate_info(cv_text, filename)
        result = calculate_score(cv_text, jd_skills)

        score  = result["score"]
        status = "Not Qualified"
        if score >= high_threshold:
            status = "Qualified"
        elif score >= low_threshold:
            status = "Review"

        if groq_key:
            questions = await generate_questions(final_jd, info["name"], result["matched"], groq_key, groq_model)
        else:
            questions = _default_questions(info["name"], result["matched"])

        resume_summary = await generate_resume_summary(cv_text, groq_key, groq_model)

        # Categorized strengths breakdown — same schema as CVAnalysis and
        # JobHunter (utils/llm_extraction.py), evaluated against this
        # session's already-extracted JD requirements so "Essential Matched"
        # reflects this specific JD, not a generic skill dump.
        from utils.llm_extraction import extract_candidate_strengths
        strengths_breakdown = await extract_candidate_strengths(
            cv_text,
            {"essential": jd_details.get("essential", []), "good_to_have": jd_details.get("good_to_have", [])},
            groq_key, groq_model,
        )

        fname_lower = (filename or "").lower()
        if fname_lower.endswith(".pdf"):
            resume_mimetype = "application/pdf"
        elif fname_lower.endswith(".docx"):
            resume_mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif fname_lower.endswith(".doc"):
            resume_mimetype = "application/msword"
        else:
            resume_mimetype = "text/plain"

        return JobLensCandidate(
            session_id=session.id,
            name=info["name"],
            email=info["email"],
            phone=info["phone"],
            filename=filename,
            ats_score=score,
            status=status,
            matched_skills=result["matched"],
            missing_skills=result["gap"],
            bonus=result["bonus"],
            bonus_reasons=result["reasons"],
            experience_years=info.get("experience_years", ""),
            summary=info.get("summary", ""),
            resume_summary=resume_summary,
            interview_questions=questions,
            video_status="Pending",
            shortlisted=False,
            resume_file_blob=content,
            resume_file_mimetype=resume_mimetype,
            source_vendor_id=source_vendor_id,
            source_vendor_name=source_vendor_name,
            source_tracked_candidate_id=source_tracked_candidate_id,
            strengths_breakdown={
                "essentialMatched": strengths_breakdown.get("essential_matched", []),
                "technicalSkills": strengths_breakdown.get("technical_skills", []),
                "businessSkills": strengths_breakdown.get("business_skills", []),
                "softSkills": strengths_breakdown.get("soft_skills", []),
                "significantExperience": strengths_breakdown.get("significant_experience", []),
                "certificationsDegrees": strengths_breakdown.get("certifications_degrees", []),
            },
        )

    candidates = []
    for upload in cv_files:
        try:
            content = await upload.read()
            candidate = await _score_and_build_candidate(content, upload.filename)
            if candidate:
                db.add(candidate)
                candidates.append(candidate)
        except Exception as e:
            print(f"Error processing {upload.filename}: {e}")
            continue

    # ── Score candidates sourced from Vendor Management submissions ──────
    for tc in source_candidates:
        try:
            if not tc.resume_blob:
                continue
            vendor_row = await db.execute(select(VendorModel).where(VendorModel.id == tc.vendor_id))
            vendor = vendor_row.scalar_one_or_none()
            candidate = await _score_and_build_candidate(
                tc.resume_blob, tc.resume_filename or f"{tc.name}.pdf",
                source_vendor_id=tc.vendor_id,
                source_vendor_name=vendor.name if vendor else "",
                source_tracked_candidate_id=tc.id,
            )
            if candidate:
                # Prefer the vendor-submitted contact details if the resume
                # parse didn't find them
                candidate.name = tc.name or candidate.name
                candidate.email = candidate.email or tc.email or ""
                candidate.phone = candidate.phone or tc.phone or ""
                db.add(candidate)
                candidates.append(candidate)
        except Exception as e:
            print(f"Error processing vendor candidate {tc.id}: {e}")
            continue

    await db.commit()
    for c in candidates:
        await db.refresh(c)

    candidates.sort(key=lambda c: c.ats_score, reverse=True)

    return {
        "session_id": session.id,
        "jd_skills":  jd_skills[:30],
        "ai_powered": groq_key is not None,
        "total":      len(candidates),
        "candidates": [_fmt(c) for c in candidates],
    }


@router.get("/sessions")
async def list_sessions(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(JobLensSession)
        .where(JobLensSession.user_id == current_user.id)
        .order_by(JobLensSession.created_at.desc())
    )
    return [
        {
            "id": s.id, "sequence_number": s.sequence_number or s.id, "cv_count": s.cv_count,
            "low_threshold": s.low_threshold, "high_threshold": s.high_threshold,
            "status": s.status,
            "created_at": s.created_at.isoformat() if s.created_at else None,
            "jd_preview": (s.jd_text or "")[:120] + "...",
            "ai_powered": False,
        }
        for s in r.scalars().all()
    ]


@router.get("/sessions/{session_id}")
async def get_session(
    session_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    sr = await db.execute(
        select(JobLensSession).where(
            JobLensSession.id == session_id,
            JobLensSession.user_id == current_user.id,
        )
    )
    session = sr.scalar_one_or_none()
    if not session:
        raise HTTPException(404, "Session not found")

    cr = await db.execute(
        select(JobLensCandidate)
        .where(JobLensCandidate.session_id == session_id)
        .order_by(JobLensCandidate.ats_score.desc())
    )
    candidates = cr.scalars().all()

    return {
        "id": session.id,
        "sequence_number": session.sequence_number or session.id,
        "jd_text": session.jd_text,
        "jd_skills": session.jd_skills,
        "jd_role": session.jd_role or "",
        "jd_location": session.jd_location or "",
        "jd_company": session.jd_company or "",
        "jd_record_id": session.jd_record_id,
        "jd_client_name": session.jd_client_name or session.jd_company or "",
        "jd_essential_skills": session.jd_essential_skills or [],
        "jd_good_to_have_skills": session.jd_good_to_have_skills or [],
        "jd_optional_skills": session.jd_optional_skills or [],
        "low_threshold": session.low_threshold,
        "high_threshold": session.high_threshold,
        "status": session.status,
        "cv_count": session.cv_count,
        "created_at": session.created_at.isoformat() if session.created_at else None,
        "candidates": [_fmt(c) for c in candidates],
    }


@router.post("/sessions/{session_id}/candidates/{candidate_id}/questions")
async def get_questions(
    session_id: int,
    candidate_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    sr = await db.execute(
        select(JobLensSession).where(
            JobLensSession.id == session_id,
            JobLensSession.user_id == current_user.id,
        )
    )
    session = sr.scalar_one_or_none()
    if not session:
        raise HTTPException(404, "Session not found")

    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.id == candidate_id)
    )
    candidate = cr.scalar_one_or_none()
    if not candidate:
        raise HTTPException(404, "Candidate not found")

    # Return existing questions or regenerate
    if candidate.interview_questions:
        return {"questions": candidate.interview_questions, "ai_powered": False}

    groq_key = await get_credential(db, current_user.id, "groq", "api_key")
    groq_model = await get_groq_model(db, current_user.id)

    if groq_key:
        questions = await generate_questions(
            session.jd_text or "", candidate.name,
            candidate.matched_skills or [], groq_key, groq_model
        )
    else:
        questions = _default_questions(candidate.name, candidate.matched_skills or [])

    candidate.interview_questions = questions
    await db.commit()
    return {"questions": questions, "ai_powered": groq_key is not None}


@router.put("/candidates/{candidate_id}/shortlist")
async def toggle_shortlist(
    candidate_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.id == candidate_id)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")
    c.shortlisted = not c.shortlisted
    await db.commit()
    return {"shortlisted": c.shortlisted}


# ── CANDIDATE CONTACT / VIDEO INTERVIEW INVITE ───────────────────────────────

@router.post("/candidates/{candidate_id}/prepare-invite")
async def prepare_invite(
    candidate_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Ensures the candidate has a unique interview token and returns it so
    the frontend can build a shareable, no-login video-interview link."""
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.id == candidate_id)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")

    if not c.interview_token:
        c.interview_token = secrets.token_urlsafe(24)
        await db.commit()
        await db.refresh(c)

    return {
        "token": c.interview_token,
        "candidate_name": c.name,
        "candidate_email": c.email,
    }


@router.post("/candidates/{candidate_id}/send-invite")
async def send_invite(
    candidate_id: int,
    payload: SendInviteRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.id == candidate_id)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")

    smtp_cfg = await _get_smtp_config(current_user.id, db)
    _send_email(smtp_cfg, payload.to_email, payload.subject, payload.body_html)

    return {"sent": True}


@router.post("/candidates/{candidate_id}/mark-contacted")
async def mark_contacted(
    candidate_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Flips the 'contacted' flag once the recruiter sends the interview
    invite letter (via mailto handoff to their own mail client)."""
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.id == candidate_id)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")
    c.contacted = True
    await db.commit()
    return {"contacted": True}


@router.get("/morphcast-key")
async def get_morphcast_key(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Returns the recruiter's saved MorphCast license key (Settings > API
    Keys, service: morphcast, key: license_key), used client-side by the
    video interview modal for facial emotion analysis. Not a sensitive
    secret in the traditional sense — MorphCast's SDK runs entirely in the
    browser and the key is visible in that SDK's own network calls anyway."""
    key = await get_credential(db, current_user.id, "morphcast", "license_key")
    return {"license_key": key or ""}


class InterviewResult(BaseModel):
    happy: int = 0
    neutral: int = 0
    sad: int = 0
    angry: int = 0
    fear: int = 0
    disgust: int = 0
    surprise: int = 0
    dominant: str = "Neutral"


@router.post("/candidates/{candidate_id}/interview-result")
async def save_interview_result(
    candidate_id: int,
    result: InterviewResult,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.id == candidate_id)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")
    c.emotion_happy    = result.happy
    c.emotion_neutral  = result.neutral
    c.emotion_sad      = result.sad
    c.emotion_angry    = result.angry
    c.emotion_fear     = result.fear
    c.emotion_disgust  = result.disgust
    c.emotion_surprise = result.surprise
    c.dominant_emotion = result.dominant
    c.video_status     = "Completed"
    await db.commit()
    return {"status": "saved"}


# ══════════════════════════════════════════════════════════════════════════════
# AUTOMATIC VIDEO ANALYSIS — runs once the video blob is stored.
# Transcribes via Groq's Whisper endpoint (accepts .webm directly, no ffmpeg
# needed), then scores the transcript against the interview questions with
# an LLM. Result is written back onto the SAME candidate row. Runs as a
# FastAPI background task so the upload request returns immediately rather
# than blocking on transcription + LLM latency.
# ══════════════════════════════════════════════════════════════════════════════

def _transcribe_video(video_bytes: bytes, mimetype: str, groq_key: str) -> str:
    """Groq's /audio/transcriptions endpoint is OpenAI-Whisper-API-compatible
    and accepts webm/mp4/mp3/wav/m4a/ogg directly — no local audio
    extraction/conversion needed."""
    resp = requests.post(
        "https://api.groq.com/openai/v1/audio/transcriptions",
        headers={"Authorization": f"Bearer {groq_key}"},
        files={"file": ("interview.webm", video_bytes, mimetype or "video/webm")},
        data={"model": "whisper-large-v3", "response_format": "text"},
        timeout=180,
        proxies={"http": None, "https": None},  # avoid any system proxy intercepting this call
    )
    resp.raise_for_status()
    return resp.text.strip()


async def _analyze_transcript(
    transcript: str, questions: list, candidate_name: str, groq_key: str, groq_model: str = DEFAULT_GROQ_MODEL
) -> dict:
    from langchain_groq import ChatGroq
    from langchain.schema import HumanMessage

    llm = ChatGroq(api_key=groq_key, model=groq_model, temperature=0.2)
    questions_block = "\n".join(f"{i+1}. {q}" for i, q in enumerate(questions)) or "(not recorded)"
    prompt = f"""You are an experienced hiring manager reviewing a recorded video
interview transcript for {candidate_name}. Be fair and evidence-based — only
comment on what the transcript actually supports.

QUESTIONS ASKED:
{questions_block}

TRANSCRIPT (auto-generated, may contain minor recognition errors):
\"\"\"{transcript[:6000]}\"\"\"

Assess the candidate's spoken interview performance. Return ONLY valid JSON,
no markdown, no commentary:
{{
  "communication_score": <0-100, clarity/structure of spoken answers>,
  "relevance_score": <0-100, how directly answers addressed the questions asked>,
  "confidence_score": <0-100, based on language used — decisiveness, specificity, hedging>,
  "overall_score": <0-100, holistic>,
  "strengths": ["<specific, evidence-based>", "..."],
  "concerns": ["<specific, evidence-based>", "..."],
  "key_observations": ["<notable moment or answer>", "..."],
  "summary": "<3-4 sentence overall assessment>"
}}"""
    resp = llm.invoke([HumanMessage(content=prompt)])
    raw = resp.content.strip()
    raw = re.sub(r"```(?:json)?|```", "", raw).strip()
    return json.loads(raw)


async def _run_video_analysis(candidate_id: int):
    """Background task — opens its OWN DB session since the request-scoped
    one is already closed by the time this runs after the response returns."""
    async with AsyncSessionLocal() as db:
        try:
            cr = await db.execute(
                select(JobLensCandidate, JobLensSession)
                .join(JobLensSession, JobLensCandidate.session_id == JobLensSession.id)
                .where(JobLensCandidate.id == candidate_id)
            )
            row = cr.first()
            if not row:
                return
            c, session = row

            c.video_analysis_status = "Processing"
            await db.commit()

            groq_key = await get_credential(db, session.user_id, "groq", "api_key")
            groq_model = await get_groq_model(db, session.user_id)
            if not groq_key:
                c.video_analysis_status = "Failed"
                c.video_analysis = {"error": "No Groq API key configured (own or admin-shared) — required for transcription and analysis."}
                await db.commit()
                return

            if not c.video_blob:
                c.video_analysis_status = "Failed"
                c.video_analysis = {"error": "No video stored for this candidate."}
                await db.commit()
                return

            transcript = _transcribe_video(c.video_blob, c.video_mimetype, groq_key)
            if not transcript:
                c.video_analysis_status = "Failed"
                c.video_analysis = {"error": "Transcription returned no speech content."}
                await db.commit()
                return

            analysis = await _analyze_transcript(
                transcript, c.interview_questions or [], c.name or "the candidate", groq_key, groq_model
            )

            c.video_transcript = transcript
            c.video_analysis = analysis
            c.video_analysis_status = "Completed"
            await db.commit()
        except Exception as e:
            try:
                c.video_analysis_status = "Failed"
                c.video_analysis = {"error": str(e)[:300]}
                await db.commit()
            except Exception:
                pass


@router.post("/candidates/{candidate_id}/video")
async def upload_interview_video(
    candidate_id: int,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Stores the recorded interview video as a blob on the candidate row,
    alongside their resume — then kicks off automatic transcription +
    performance analysis in the background, written back onto this same row."""
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.id == candidate_id)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")
    content = await file.read()
    c.video_blob = content
    c.video_mimetype = file.content_type or "video/webm"
    c.video_analysis_status = "Pending"
    await db.commit()
    background_tasks.add_task(_run_video_analysis, candidate_id)
    return {"status": "saved", "size_bytes": len(content)}


@router.post("/candidates/{candidate_id}/reanalyze-video")
async def reanalyze_video(
    candidate_id: int,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Re-runs transcription + performance analysis on the ALREADY-STORED
    video for this candidate — no re-recording or re-upload needed. Useful
    after adding a Groq key, or to re-check with updated interview
    questions/context."""
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.id == candidate_id)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")
    if not c.video_blob:
        raise HTTPException(400, "No video stored for this candidate yet.")
    c.video_analysis_status = "Pending"
    await db.commit()
    background_tasks.add_task(_run_video_analysis, candidate_id)
    return {"status": "queued"}
async def download_interview_video(
    candidate_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    cr = await db.execute(
        select(JobLensCandidate, JobLensSession)
        .join(JobLensSession, JobLensCandidate.session_id == JobLensSession.id)
        .where(JobLensCandidate.id == candidate_id)
    )
    row = cr.first()
    if not row:
        raise HTTPException(404, "Candidate not found")
    c, session = row
    # Row-level access control: only the owning recruiter or an admin may
    # view this candidate's video — never any other user.
    if session.user_id != current_user.id and current_user.role != "admin":
        raise HTTPException(404, "Candidate not found")
    if not c.video_blob:
        raise HTTPException(404, "No video recorded for this candidate")
    return Response(content=c.video_blob, media_type=c.video_mimetype or "video/webm")


@router.get("/candidates/{candidate_id}/resume-file")
async def download_resume_file(
    candidate_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    cr = await db.execute(
        select(JobLensCandidate, JobLensSession)
        .join(JobLensSession, JobLensCandidate.session_id == JobLensSession.id)
        .where(JobLensCandidate.id == candidate_id)
    )
    row = cr.first()
    if not row:
        raise HTTPException(404, "Candidate not found")
    c, session = row
    if session.user_id != current_user.id and current_user.role != "admin":
        raise HTTPException(404, "Candidate not found")
    if not c.resume_file_blob:
        raise HTTPException(404, "No resume file stored for this candidate")
    return Response(
        content=c.resume_file_blob,
        media_type=c.resume_file_mimetype or "application/octet-stream",
        headers={"Content-Disposition": f'inline; filename="{c.filename or "resume"}"'},
    )
# These power the emailed video-interview link so a candidate can complete
# the interview without a TalentIQ login. Access is gated by the unguessable
# token, not a session/auth check.

@router.get("/public/interview/{token}")
async def public_get_interview(token: str, db: AsyncSession = Depends(get_db)):
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.interview_token == token)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "This interview link is invalid or has expired.")
    return {
        "candidate_name": c.name,
        "questions": c.interview_questions or [],
        "video_status": c.video_status,
    }


@router.get("/public/interview/{token}/morphcast-key")
async def public_get_morphcast_key(token: str, db: AsyncSession = Depends(get_db)):
    """Same MorphCast license key as the authenticated endpoint above, but
    resolved via the interview token (no login) — belongs to whichever
    recruiter's account generated this candidate's interview."""
    cr = await db.execute(
        select(JobLensCandidate, JobLensSession)
        .join(JobLensSession, JobLensCandidate.session_id == JobLensSession.id)
        .where(JobLensCandidate.interview_token == token)
    )
    row = cr.first()
    if not row:
        raise HTTPException(404, "This interview link is invalid or has expired.")
    _, session = row

    key = await get_credential(db, session.user_id, "morphcast", "license_key")
    return {"license_key": key or ""}


@router.post("/public/interview/{token}/result")
async def public_save_interview_result(
    token: str,
    result: InterviewResult,
    db: AsyncSession = Depends(get_db),
):
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.interview_token == token)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "This interview link is invalid or has expired.")
    c.emotion_happy    = result.happy
    c.emotion_neutral  = result.neutral
    c.emotion_sad      = result.sad
    c.emotion_angry    = result.angry
    c.emotion_fear     = result.fear
    c.emotion_disgust  = result.disgust
    c.emotion_surprise = result.surprise
    c.dominant_emotion = result.dominant
    c.video_status     = "Completed"
    await db.commit()
    return {"status": "saved"}


@router.post("/public/interview/{token}/video")
async def public_upload_interview_video(
    token: str,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    """Same as the authenticated video-upload endpoint, but reached via the
    interview token so the candidate (no login) can submit their recorded
    video directly from the public interview page."""
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.interview_token == token)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "This interview link is invalid or has expired.")
    content = await file.read()
    c.video_blob = content
    c.video_mimetype = file.content_type or "video/webm"
    c.video_analysis_status = "Pending"
    await db.commit()
    background_tasks.add_task(_run_video_analysis, c.id)
    return {"status": "saved", "size_bytes": len(content)}


@router.get("/sessions/{session_id}/export")
async def export_session(
    session_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    import pandas as pd

    sr = await db.execute(
        select(JobLensSession).where(
            JobLensSession.id == session_id,
            JobLensSession.user_id == current_user.id,
        )
    )
    session = sr.scalar_one_or_none()
    if not session:
        raise HTTPException(404, "Session not found")

    cr = await db.execute(
        select(JobLensCandidate)
        .where(JobLensCandidate.session_id == session_id)
        .order_by(JobLensCandidate.ats_score.desc())
    )
    candidates = cr.scalars().all()

    rows = []
    for i, c in enumerate(candidates, 1):
        rows.append({
            "Rank":              i,
            "Name":              c.name,
            "Email":             c.email,
            "Phone":             c.phone,
            "Resume Summary":    " | ".join(c.resume_summary or []),
            "ATS Score":         f"{c.ats_score:.1f}%",
            "Key Strength":      ", ".join(c.matched_skills or []),
            "Considerations":    ", ".join(c.missing_skills or []),
            "Status":            c.status,
            "Video Status":      c.video_status,
            "Happy %":           c.emotion_happy or 0,
            "Neutral %":         c.emotion_neutral or 0,
            "Sad %":             c.emotion_sad or 0,
            "Angry %":           c.emotion_angry or 0,
            "Fear %":            c.emotion_fear or 0,
            "Disgust %":         c.emotion_disgust or 0,
            "Surprise %":        c.emotion_surprise or 0,
            "Dominant Emotion":  c.dominant_emotion or "Neutral",
            "Shortlisted":       "Yes" if c.shortlisted else "No",
            "Bonus Points":      c.bonus,
            "Bonus Reasons":     c.bonus_reasons,
            "Summary":           c.summary or "",
        })

    buf = io.BytesIO()
    pd.DataFrame(rows).to_excel(buf, index=False)
    buf.seek(0)

    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=candidatelens_{session_id}.xlsx"},
    )


@router.delete("/sessions/{session_id}")
async def delete_session(
    session_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Delete a CandidateLens session and all its candidates."""
    from sqlalchemy import delete as sql_delete
    result = await db.execute(
        select(JobLensSession).where(
            JobLensSession.id == session_id,
            JobLensSession.user_id == current_user.id,
        )
    )
    session = result.scalar_one_or_none()
    if not session:
        raise HTTPException(404, "Session not found")
    await db.execute(sql_delete(JobLensCandidate).where(JobLensCandidate.session_id == session_id))
    await db.delete(session)
    await db.commit()
    return {"message": "Deleted"}


@router.delete("/sessions")
async def delete_all_sessions(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Delete ALL sessions for the current user."""
    from sqlalchemy import delete as sql_delete
    ids_r = await db.execute(
        select(JobLensSession.id).where(JobLensSession.user_id == current_user.id)
    )
    ids = [r[0] for r in ids_r.all()]
    if ids:
        await db.execute(sql_delete(JobLensCandidate).where(JobLensCandidate.session_id.in_(ids)))
        await db.execute(sql_delete(JobLensSession).where(JobLensSession.user_id == current_user.id))
        await db.commit()
    return {"message": f"Deleted {len(ids)} sessions"}